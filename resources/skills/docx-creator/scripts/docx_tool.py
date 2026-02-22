#!/usr/bin/env python3
"""
Create, read, and edit Word documents (.docx).
Dependencies: python-docx
"""

import argparse
import csv
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt
    pass  # docx.enum.text available if needed
except ImportError:
    print("Missing dependency: python-docx", file=sys.stderr)
    print("Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Markdown → DOCX conversion helpers
# ---------------------------------------------------------------------------

def parse_markdown_to_docx(doc, md_text):
    """Convert Markdown text to docx paragraphs."""
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line):
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            i += 1
            continue

        # Table (pipe-delimited)
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[-:|]+\|', lines[i + 1]):
            rows = []
            while i < len(lines) and '|' in lines[i]:
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                # Skip separator line
                if not re.match(r'^[-:|]+$', cells[0] if cells else ''):
                    rows.append(cells)
                i += 1
            if rows:
                num_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            table.rows[ri].cells[ci].text = cell_text
            continue

        # Unordered list
        list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if list_match:
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, text)
            i += 1
            continue

        # Ordered list
        olist_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if olist_match:
            text = olist_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, text)
            i += 1
            continue

        # Code block
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, line)
        i += 1


def add_formatted_run(paragraph, text):
    """Add text to paragraph with basic inline formatting (bold, italic, code)."""
    # Simple inline formatting: **bold**, *italic*, `code`
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif part:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------

def cmd_create(args):
    """Create a new .docx document."""
    if args.template and os.path.isfile(args.template):
        doc = Document(args.template)
    else:
        doc = Document()

    # Set metadata
    if args.author:
        doc.core_properties.author = args.author

    # Add title
    if args.title:
        doc.add_heading(args.title, level=0)

    # Add content from Markdown file
    if args.from_markdown:
        if not os.path.isfile(args.from_markdown):
            print(f"Error: File not found: {args.from_markdown}", file=sys.stderr)
            sys.exit(1)
        with open(args.from_markdown, encoding='utf-8') as f:
            md_text = f.read()
        parse_markdown_to_docx(doc, md_text)
    elif args.content:
        doc.add_paragraph(args.content)

    doc.save(args.output)
    print(f"Created: {args.output}")


def cmd_read(args):
    """Read / extract text from a .docx document."""
    if not os.path.isfile(args.input):
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    doc = Document(args.input)
    fmt = args.format

    if fmt == 'json':
        paragraphs = []
        for p in doc.paragraphs:
            paragraphs.append({
                'style': p.style.name if p.style else 'Normal',
                'text': p.text,
                'alignment': str(p.alignment) if p.alignment else None
            })
        output = json.dumps(paragraphs, indent=2, ensure_ascii=False)

    elif fmt == 'markdown':
        lines = []
        for p in doc.paragraphs:
            style = p.style.name if p.style else ''
            text = p.text
            if 'Heading 1' in style:
                lines.append(f'# {text}')
            elif 'Heading 2' in style:
                lines.append(f'## {text}')
            elif 'Heading 3' in style:
                lines.append(f'### {text}')
            elif 'Heading' in style:
                level = 4
                m = re.search(r'(\d+)', style)
                if m:
                    level = int(m.group(1))
                lines.append(f'{"#" * level} {text}')
            elif 'List Bullet' in style:
                lines.append(f'- {text}')
            elif 'List Number' in style:
                lines.append(f'1. {text}')
            elif text.strip():
                lines.append(text)
            else:
                lines.append('')
        # Also extract tables
        for table in doc.tables:
            lines.append('')
            for ri, row in enumerate(table.rows):
                cells = [cell.text for cell in row.cells]
                lines.append('| ' + ' | '.join(cells) + ' |')
                if ri == 0:
                    lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
            lines.append('')
        output = '\n'.join(lines)

    else:  # text
        lines = [p.text for p in doc.paragraphs]
        # Include table text
        for table in doc.tables:
            for row in table.rows:
                lines.append('\t'.join(cell.text for cell in row.cells))
        output = '\n'.join(lines)

    if args.save:
        with open(args.save, 'w', encoding='utf-8') as f:
            f.write(output)
        print(f"Saved to: {args.save}", file=sys.stderr)
    else:
        print(output)


def cmd_append(args):
    """Append content to an existing .docx document."""
    if not os.path.isfile(args.input):
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    doc = Document(args.input)

    if args.heading:
        doc.add_heading(args.heading, level=args.level)
        print(f"Added heading (level {args.level}): {args.heading}")

    if args.paragraph:
        p = doc.add_paragraph()
        add_formatted_run(p, args.paragraph)
        print("Added paragraph")

    if args.table:
        cols = [c.strip() for c in args.table.split(',')]
        num_cols = len(cols)
        rows_data = []
        if args.rows:
            for row_str in args.rows.split('|'):
                rows_data.append([c.strip() for c in row_str.split(',')])

        table = doc.add_table(rows=1 + len(rows_data), cols=num_cols, style='Table Grid')
        for ci, col_name in enumerate(cols):
            table.rows[0].cells[ci].text = col_name
        for ri, row_data in enumerate(rows_data):
            for ci, cell_text in enumerate(row_data):
                if ci < num_cols:
                    table.rows[ri + 1].cells[ci].text = cell_text
        print(f"Added table ({len(rows_data)} rows x {num_cols} cols)")

    if args.image:
        if os.path.isfile(args.image):
            width = Inches(args.image_width) if args.image_width else Inches(4)
            doc.add_picture(args.image, width=width)
            print(f"Added image: {args.image}")
        else:
            print(f"Warning: Image not found: {args.image}", file=sys.stderr)

    doc.save(args.input)
    print(f"Updated: {args.input}")


def cmd_table(args):
    """Add a table from CSV data to a document."""
    if not os.path.isfile(args.csv_file):
        print(f"Error: CSV file not found: {args.csv_file}", file=sys.stderr)
        sys.exit(1)

    with open(args.csv_file, newline='', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        print("CSV file is empty.", file=sys.stderr)
        sys.exit(1)

    if os.path.isfile(args.output):
        doc = Document(args.output)
    else:
        doc = Document()

    if args.title:
        doc.add_heading(args.title, level=2)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Table Grid')
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < len(table.columns):
                table.rows[ri].cells[ci].text = cell_text

    # Bold the header row
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.save(args.output)
    print(f"Added table ({len(rows)-1} data rows) to {args.output}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description='Word Document Tool')
    sub = parser.add_subparsers(dest='command', required=True)

    # create
    p = sub.add_parser('create', help='Create a new .docx document')
    p.add_argument('output', help='Output .docx file path')
    p.add_argument('--title', help='Document title')
    p.add_argument('--author', help='Document author')
    p.add_argument('--from-markdown', help='Create from Markdown file')
    p.add_argument('--content', help='Inline body text')
    p.add_argument('--template', help='Template .docx for styles')

    # read
    p = sub.add_parser('read', help='Read/extract text from .docx')
    p.add_argument('input', help='Input .docx file path')
    p.add_argument('--format', default='text', choices=['text', 'markdown', 'json'])
    p.add_argument('--save', help='Save output to file')

    # append
    p = sub.add_parser('append', help='Append content to existing .docx')
    p.add_argument('input', help='.docx file to modify')
    p.add_argument('--heading', help='Add a heading')
    p.add_argument('--level', type=int, default=2, help='Heading level (default: 2)')
    p.add_argument('--paragraph', help='Add a paragraph')
    p.add_argument('--table', help='Add table with comma-separated column names')
    p.add_argument('--rows', help='Table row data: "a,b,c|d,e,f" (pipe-separated rows)')
    p.add_argument('--image', help='Add an image file')
    p.add_argument('--image-width', type=float, help='Image width in inches (default: 4)')

    # table
    p = sub.add_parser('table', help='Add table from CSV to .docx')
    p.add_argument('output', help='.docx file (created or appended)')
    p.add_argument('--from-csv', dest='csv_file', required=True, help='CSV file path')
    p.add_argument('--title', help='Table heading')

    args = parser.parse_args()

    commands = {
        'create': cmd_create,
        'read': cmd_read,
        'append': cmd_append,
        'table': cmd_table,
    }
    commands[args.command](args)


if __name__ == '__main__':
    main()
